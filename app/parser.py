"""Resume Document Parser
Extracts text from PDF, DOCX, TXT, and raw string formats.
"""

import io
import os
from typing import Union, BinaryIO


class ResumeParser:
    """Handles extracting clean text from various resume document formats."""

    @staticmethod
    def extract_from_txt(source: Union[str, BinaryIO, bytes]) -> str:
        """Extract text from plain text file path, bytes, or file stream."""
        if isinstance(source, str):
            if os.path.exists(source):
                with open(source, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
                return source  # Return direct text if not a file path
            return source
        elif isinstance(source, bytes):
            return source.decode("utf-8", errors="ignore")
        elif hasattr(source, "read"):
            content = source.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="ignore")
            return str(content)
        return str(source)

    @staticmethod
    def extract_from_pdf(source: Union[str, BinaryIO, bytes]) -> str:
        """Extract text from a PDF file path or stream using pypdf/PyPDF2."""
        text_content = []
        try:
            try:
                import pypdf as pdf_module
            except ImportError:
                import PyPDF2 as pdf_module  # type: ignore

            if isinstance(source, str) and os.path.exists(source):
                reader = pdf_module.PdfReader(source)
            elif isinstance(source, bytes):
                reader = pdf_module.PdfReader(io.BytesIO(source))
            elif hasattr(source, "read"):
                reader = pdf_module.PdfReader(source)
            else:
                raise ValueError("Invalid PDF source provided.")

            for page_idx, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text_content.append(extracted)
            return "\n".join(text_content).strip()
        except ImportError:
            raise ImportError(
                "PDF parsing requires 'pypdf' or 'PyPDF2'. Please run: pip install pypdf"
            )
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF file: {str(e)}")

    @staticmethod
    def extract_from_docx(source: Union[str, BinaryIO, bytes]) -> str:
        """Extract text from a DOCX file path or stream."""
        try:
            import docx

            if isinstance(source, str) and os.path.exists(source):
                doc = docx.Document(source)
            elif isinstance(source, bytes):
                doc = docx.Document(io.BytesIO(source))
            elif hasattr(source, "read"):
                doc = docx.Document(source)
            else:
                raise ValueError("Invalid DOCX source provided.")

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs).strip()
        except ImportError:
            raise ImportError(
                "DOCX parsing requires 'python-docx'. Please run: pip install python-docx"
            )
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX file: {str(e)}")

    @classmethod
    def parse(cls, source: Union[str, BinaryIO, bytes], filename: str = "") -> str:
        """Auto-detect format and parse resume into clean string."""
        if isinstance(source, str):
            # Check if source is a file path
            if os.path.isfile(source):
                ext = os.path.splitext(source)[1].lower()
                if ext == ".pdf":
                    return cls.extract_from_pdf(source)
                elif ext in [".docx", ".doc"]:
                    return cls.extract_from_docx(source)
                else:
                    return cls.extract_from_txt(source)
            # If not a file path, treat as raw text
            return source.strip()

        # Handle streams with filename hint
        if filename:
            ext = os.path.splitext(filename)[1].lower()
            if ext == ".pdf":
                return cls.extract_from_pdf(source)
            elif ext in [".docx", ".doc"]:
                return cls.extract_from_docx(source)
            else:
                return cls.extract_from_txt(source)

        # Fallback to plain text extraction
        return cls.extract_from_txt(source)

